import os
import tempfile
from typing import Dict, Any, Optional, Tuple
from fastapi import UploadFile, HTTPException
import PyPDF2
from docx import Document
import pandas as pd
from datetime import datetime

class DocumentService:
    SUPPORTED_EXTENSIONS = {
        '.pdf': 'application/pdf',
        '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        '.csv': 'text/csv',
        '.txt': 'text/plain'
    }

    def __init__(self):
        self.temp_dir = tempfile.mkdtemp()

    def validate_file(self, file: UploadFile) -> Tuple[str, str]:
        """���� ��ȿ�� �˻� �� Ȯ���� Ȯ��"""
        ext = os.path.splitext(file.filename)[1].lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise HTTPException(
                status_code=400,
                detail=f"�������� �ʴ� ���� �����Դϴ�. ���� ����: {', '.join(self.SUPPORTED_EXTENSIONS.keys())}"
            )
        return ext, self.SUPPORTED_EXTENSIONS[ext]

    async def save_upload_file(self, file: UploadFile) -> str:
        """���ε�� ������ �ӽ� ����"""
        try:
            ext, _ = self.validate_file(file)
            temp_path = os.path.join(self.temp_dir, f"{datetime.now().strftime('%Y%m%d_%H%M%S')}{ext}")
            
            with open(temp_path, "wb") as buffer:
                content = await file.read()
                buffer.write(content)
            
            return temp_path
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"���� ���� ����: {str(e)}")

    def extract_text_from_pdf(self, file_path: str) -> Dict[str, Any]:
        """PDF ���Ͽ��� �ؽ�Ʈ ����"""
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                
                return {
                    "text": text,
                    "pages": len(reader.pages),
                    "word_count": len(text.split())
                }
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"PDF �ؽ�Ʈ ���� ����: {str(e)}")

    def extract_text_from_docx(self, file_path: str) -> Dict[str, Any]:
        """DOCX ���Ͽ��� �ؽ�Ʈ ����"""
        try:
            doc = Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            return {
                "text": text,
                "paragraphs": len(doc.paragraphs),
                "word_count": len(text.split())
            }
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"DOCX �ؽ�Ʈ ���� ����: {str(e)}")

    def extract_text_from_excel(self, file_path: str) -> Dict[str, Any]:
        """Excel ���Ͽ��� �ؽ�Ʈ ����"""
        try:
            df = pd.read_excel(file_path)
            text = df.to_string()
            
            return {
                "text": text,
                "rows": len(df),
                "columns": len(df.columns),
                "word_count": len(text.split())
            }
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Excel �ؽ�Ʈ ���� ����: {str(e)}")

    def extract_text_from_csv(self, file_path: str) -> Dict[str, Any]:
        """CSV ���Ͽ��� �ؽ�Ʈ ����"""
        try:
            df = pd.read_csv(file_path)
            text = df.to_string()
            
            return {
                "text": text,
                "rows": len(df),
                "columns": len(df.columns),
                "word_count": len(text.split())
            }
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"CSV �ؽ�Ʈ ���� ����: {str(e)}")

    def extract_text_from_txt(self, file_path: str) -> Dict[str, Any]:
        """TXT ���Ͽ��� �ؽ�Ʈ ����"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            return {
                "text": text,
                "lines": len(text.splitlines()),
                "word_count": len(text.split())
            }
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"TXT �ؽ�Ʈ ���� ����: {str(e)}")

    async def process_document(self, file: UploadFile) -> Dict[str, Any]:
        """���� ó�� �� �ؽ�Ʈ ����"""
        try:
            # ���� ����
            file_path = await self.save_upload_file(file)
            ext = os.path.splitext(file_path)[1].lower()
            
            # ���� ũ��
            file_size = os.path.getsize(file_path)
            
            # ���� ���ĺ� �ؽ�Ʈ ����
            if ext == '.pdf':
                result = self.extract_text_from_pdf(file_path)
            elif ext == '.docx':
                result = self.extract_text_from_docx(file_path)
            elif ext == '.xlsx':
                result = self.extract_text_from_excel(file_path)
            elif ext == '.csv':
                result = self.extract_text_from_csv(file_path)
            elif ext == '.txt':
                result = self.extract_text_from_txt(file_path)
            else:
                raise HTTPException(status_code=400, detail="�������� �ʴ� ���� �����Դϴ�")
            
            # �ӽ� ���� ����
            os.remove(file_path)
            
            return {
                "filename": file.filename,
                "content_type": self.SUPPORTED_EXTENSIONS[ext],
                "size": file_size,
                **result
            }
            
        except Exception as e:
            # �ӽ� ���� ����
            if os.path.exists(file_path):
                os.remove(file_path)
            raise HTTPException(status_code=500, detail=f"���� ó�� ����: {str(e)}")

document_service = DocumentService() 